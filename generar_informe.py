"""
Genera el informe del proyecto en formato Word (.docx).
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ── Estilos base ───────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

# ── Portada ────────────────────────────────────────────────────────────────
for _ in range(6):
    doc.add_paragraph()

titulo = doc.add_paragraph()
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = titulo.add_run('INFORME DEL PROYECTO')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Sistema Hospitalario — Prototipado y Modelado de Datos')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Materia: Programación con Python\n').font.size = Pt(11)
meta.add_run('Fase 1: Estructuras de Datos Nativas\n').font.size = Pt(11)
meta.add_run('Mayo 2026').font.size = Pt(11)

doc.add_page_break()

# ── 1. Introducción ───────────────────────────────────────────────────────
doc.add_heading('1. Introducción', level=1)
doc.add_paragraph(
    'El presente informe documenta el proceso completo de desarrollo del '
    'Sistema Hospitalario, un proyecto académico cuyo objetivo es demostrar '
    'el dominio de las estructuras de datos nativas de Python (diccionarios, '
    'tuplas, conjuntos y listas) aplicadas a un contexto real de gestión '
    'clínica. El proyecto abarca desde la etapa de planteamiento y '
    'brainstorming hasta la implementación funcional completa.'
)

# ── 2. Planteamiento del problema ─────────────────────────────────────────
doc.add_heading('2. Planteamiento del Problema', level=1)
doc.add_paragraph(
    'Se requiere un sistema capaz de gestionar la información de un entorno '
    'hospitalario: pacientes, personal médico, citas, historial clínico y '
    'flujo de atención. El sistema debe:'
)
requisitos = [
    'Utilizar exclusivamente estructuras de datos nativas de Python (sin POO).',
    'Implementar colas FIFO para el flujo de pacientes en espera.',
    'Implementar pilas LIFO para el control y reversión de medicación.',
    'Usar algoritmos de ordenamiento manual (sin sort() ni sorted()).',
    'Simular listas enlazadas con diccionarios como nodos.',
    'Demostrar operaciones de conjuntos (intersección, unión, diferencia) '
    'con datos reales de alergias.',
    'Cargar datos masivamente desde archivos CSV.',
]
for r in requisitos:
    doc.add_paragraph(r, style='List Bullet')

# ── 3. Brainstorming y Diseño ─────────────────────────────────────────────
doc.add_heading('3. Brainstorming y Diseño Inicial', level=1)

doc.add_heading('3.1 Lluvia de ideas', level=2)
doc.add_paragraph(
    'Durante la fase inicial se exploraron distintas aproximaciones para '
    'modelar un sistema hospitalario. Las ideas principales fueron:'
)
ideas = [
    ('Dominio del problema', 'Se identificaron las entidades clave: pacientes, '
     'personal médico con jerarquía de acceso, citas médicas, historial '
     'clínico y alergias.'),
    ('Elección de estructuras', 'Se evaluaron las cuatro estructuras nativas '
     'de Python y se asignó cada una al caso de uso más apropiado según su '
     'complejidad algorítmica.'),
    ('Flujo de atención', 'Se decidió modelar la sala de espera como una cola '
     'FIFO (primero en llegar, primero en ser atendido) y la medicación como '
     'una pila LIFO (deshacer el último error primero).'),
    ('Persistencia vs. memoria', 'Se optó por trabajar en memoria con '
     'capacidad de persistencia JSON y exportación CSV.'),
]
for titulo_i, desc in ideas:
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    run = p.add_run(f'{titulo_i}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('3.2 Decisiones de arquitectura', level=2)
doc.add_paragraph(
    'Inicialmente se planteó una arquitectura modular con 9 archivos Python '
    'separados (config.py, utils.py, mod_pacientes.py, mod_personal.py, '
    'mod_citas.py, mod_persistencia.py, mod_exportacion_csv.py, '
    'mod_reportes.py y main.py). Esta arquitectura se diseñó con metodología '
    'top-down: descomponer el sistema en subsistemas independientes con '
    'interfaces claras entre ellos.'
)
doc.add_paragraph(
    'Posteriormente, siguiendo la guía académica del curso, se reestructuró '
    'el proyecto a un solo archivo (sistema_hospitalario.py) enfocado '
    'exclusivamente en la estructura de datos y el tratamiento de datos, '
    'sin interfaz de usuario ni menús interactivos.'
)

# ── 4. Mapeo de estructuras ───────────────────────────────────────────────
doc.add_heading('4. Mapeo de Estructuras de Datos', level=1)
doc.add_paragraph(
    'Cada estructura de Python se eligió con una justificación algorítmica '
    'concreta:'
)

table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Estructura', 'Uso en el sistema', 'Justificación']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

data = [
    ('Diccionario (dict)', 'Catálogos de pacientes y personal',
     'Búsqueda O(1) por clave (DNI o ID)'),
    ('Tupla (tuple)', 'Datos inmutables (DNI, tipo de sangre)',
     'Inmutabilidad garantizada por el lenguaje'),
    ('Conjunto (set)', 'Alergias de pacientes, permisos',
     'Operaciones de conjuntos O(1): intersección, unión, diferencia'),
    ('Lista (list)', 'Historial, cola de espera, pila de medicación',
     'Orden secuencial, append/pop O(1), simula FIFO y LIFO'),
]
for i, (est, uso, just) in enumerate(data, 1):
    table.rows[i].cells[0].text = est
    table.rows[i].cells[1].text = uso
    table.rows[i].cells[2].text = just

# ── 5. Implementación ────────────────────────────────────────────────────
doc.add_heading('5. Implementación', level=1)

doc.add_heading('5.1 Sección 1: Estructuras de Datos Básicas y Carga CSV', level=2)
doc.add_paragraph(
    'Se definieron dos diccionarios globales como bases de datos en memoria: '
    '«pacientes» (clave: DNI) y «personal» (clave: ID_Empleado con nivel '
    'de acceso jerárquico). La función registrar_paciente() crea un registro '
    'donde los datos inmutables se almacenan en una tupla y las alergias en '
    'un conjunto (set). La función cargar_pacientes_csv() lee un archivo CSV '
    'con el módulo estándar csv, separa las alergias por guion y las '
    'convierte en lista para pasarlas a registrar_paciente().'
)

doc.add_heading('5.2 Sección 2: Flujo de Atención Dinámica', level=2)
doc.add_paragraph(
    'Se implementaron dos estructuras dinámicas fundamentales:'
)
p = doc.add_paragraph()
p.style = 'List Bullet'
run = p.add_run('Cola FIFO (lista_espera): ')
run.bold = True
p.add_run(
    'Los pacientes se agregan al final con append() y se atienden desde '
    'el inicio con pop(0), garantizando el orden de llegada.'
)

p = doc.add_paragraph()
p.style = 'List Bullet'
run = p.add_run('Pila LIFO (lista_medicacion): ')
run.bold = True
p.add_run(
    'Las recetas se apilan con append() y se deshacen con pop(-1), '
    'permitiendo revertir errores en orden inverso.'
)

doc.add_paragraph(
    'Adicionalmente, la función recetar_medicamento() integra un control '
    'de acceso jerárquico: solo el personal con nivel 2 (Médico General) o '
    'superior puede prescribir medicamentos, verificando el nivel de acceso '
    'en el diccionario de personal.'
)

doc.add_heading('5.3 Sección 3: Búsqueda, Ordenamiento e Historiales', level=2)
doc.add_paragraph(
    'Esta sección demuestra tres técnicas avanzadas:'
)

p = doc.add_paragraph()
p.style = 'List Bullet'
run = p.add_run('Lista enlazada simulada: ')
run.bold = True
p.add_run(
    'El historial clínico se modela como una lista enlazada donde cada nodo '
    'es un diccionario con las claves «diagnostico» y «siguiente». La '
    'función agregar_registro_historial() recorre la cadena hasta encontrar '
    'el último nodo (siguiente == None) y enlaza el nuevo nodo al final.'
)

p = doc.add_paragraph()
p.style = 'List Bullet'
run = p.add_run('Ordenamiento Burbuja manual: ')
run.bold = True
p.add_run(
    'La función ordenar_pacientes_por_dni() extrae las claves del '
    'diccionario de pacientes y aplica Bubble Sort con intercambio de '
    'posiciones, sin usar sort() ni sorted().'
)

p = doc.add_paragraph()
p.style = 'List Bullet'
run = p.add_run('Búsqueda rápida en conjuntos: ')
run.bold = True
p.add_run(
    'La función busqueda_rapida_alergia() realiza una doble búsqueda O(1): '
    'primero localiza al paciente en el diccionario y luego verifica la '
    'pertenencia de la alergia en el set con el operador «in».'
)

# ── 6. Funciones auxiliares ───────────────────────────────────────────────
doc.add_heading('5.4 Funciones Auxiliares', level=2)
doc.add_paragraph(
    'Se implementaron funciones de apoyo que demuestran operaciones '
    'avanzadas con conjuntos:'
)
p = doc.add_paragraph()
p.style = 'List Bullet'
run = p.add_run('mostrar_historial(): ')
run.bold = True
p.add_run('Recorre la lista enlazada nodo a nodo imprimiendo cada diagnóstico.')

p = doc.add_paragraph()
p.style = 'List Bullet'
run = p.add_run('cruzar_alergias(): ')
run.bold = True
p.add_run(
    'Ejecuta cuatro operaciones de conjuntos entre dos pacientes: '
    'intersección (&), unión (|), diferencia (-) y diferencia inversa. '
    'Esto tiene aplicación médica directa para identificar alergias '
    'comunes antes de prescribir tratamientos compartidos.'
)

# ── 6. Pruebas ────────────────────────────────────────────────────────────
doc.add_heading('6. Pruebas y Validación', level=1)
doc.add_paragraph(
    'El archivo incluye una zona de pruebas (bloque if __name__ == "__main__") '
    'con 7 pruebas funcionales que verifican cada componente:'
)

table2 = doc.add_table(rows=8, cols=2)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
table2.rows[0].cells[0].text = 'Prueba'
table2.rows[0].cells[1].text = 'Verificación'
for p in table2.rows[0].cells[0].paragraphs:
    for r in p.runs:
        r.bold = True
for p in table2.rows[0].cells[1].paragraphs:
    for r in p.runs:
        r.bold = True

pruebas = [
    ('1. Registro y Estructuras', 'Dict, tupla y set creados correctamente'),
    ('2. Búsqueda de alergias', 'Operación «in» sobre set retorna True/False'),
    ('3. Cruce de alergias', 'Intersección, unión y diferencia de sets'),
    ('4. Cola de espera (FIFO)', 'append + pop(0) respeta orden de llegada'),
    ('5. Pila de medicación (LIFO)', 'append + pop(-1) deshace el último'),
    ('6. Historial (Lista enlazada)', 'Nodos enlazados se recorren en orden'),
    ('7. Ordenamiento (Burbuja)', 'DNIs ordenados sin sort()'),
]
for i, (prueba, verif) in enumerate(pruebas, 1):
    table2.rows[i].cells[0].text = prueba
    table2.rows[i].cells[1].text = verif

# ── 7. Arquitectura modular ──────────────────────────────────────────────
doc.add_heading('7. Arquitectura Modular Complementaria', level=1)
doc.add_paragraph(
    'Además del archivo principal del proyecto académico, se desarrolló una '
    'arquitectura modular completa como ejercicio adicional de ingeniería de '
    'software. Esta versión separa el sistema en módulos independientes:'
)

table3 = doc.add_table(rows=10, cols=2)
table3.style = 'Light Grid Accent 1'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
table3.rows[0].cells[0].text = 'Módulo'
table3.rows[0].cells[1].text = 'Responsabilidad'
for p in table3.rows[0].cells[0].paragraphs:
    for r in p.runs:
        r.bold = True
for p in table3.rows[0].cells[1].paragraphs:
    for r in p.runs:
        r.bold = True

modulos = [
    ('config.py', 'Constantes globales, rutas, jerarquía y permisos'),
    ('utils.py', 'Validaciones (DNI, sangre, roles) y helpers de formateo'),
    ('mod_pacientes.py', 'CRUD de pacientes, alergias con sets, historial'),
    ('mod_personal.py', 'CRUD del personal, jerarquía, control de acceso'),
    ('mod_citas.py', 'Citas, cola FIFO, pila LIFO de medicación'),
    ('mod_persistencia.py', 'Serialización JSON con soporte para sets y tuplas'),
    ('mod_exportacion_csv.py', 'Exportación de reportes a archivos CSV'),
    ('mod_reportes.py', 'Generación de reportes y estadísticas'),
    ('main.py', 'Menú interactivo por consola con 4 submenús'),
]
for i, (mod, resp) in enumerate(modulos, 1):
    table3.rows[i].cells[0].text = mod
    table3.rows[i].cells[1].text = resp

# ── 8. Complejidad ───────────────────────────────────────────────────────
doc.add_heading('8. Análisis de Complejidad Algorítmica', level=1)

table4 = doc.add_table(rows=8, cols=3)
table4.style = 'Light Grid Accent 1'
table4.alignment = WD_TABLE_ALIGNMENT.CENTER
table4.rows[0].cells[0].text = 'Operación'
table4.rows[0].cells[1].text = 'Complejidad'
table4.rows[0].cells[2].text = 'Estructura utilizada'
for cell in table4.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

complejidades = [
    ('Buscar paciente por DNI', 'O(1)', 'Diccionario'),
    ('Verificar alergia', 'O(1)', 'Set'),
    ('Agregar a cola de espera', 'O(1)', 'Lista (append)'),
    ('Atender siguiente paciente', 'O(n)', 'Lista (pop(0))'),
    ('Deshacer medicación', 'O(1)', 'Lista (pop(-1))'),
    ('Agregar al historial', 'O(n)', 'Lista enlazada (recorrido)'),
    ('Ordenar por DNI', 'O(n²)', 'Bubble Sort manual'),
]
for i, (op, comp, est) in enumerate(complejidades, 1):
    table4.rows[i].cells[0].text = op
    table4.rows[i].cells[1].text = comp
    table4.rows[i].cells[2].text = est

# ── 9. Conclusiones ──────────────────────────────────────────────────────
doc.add_heading('9. Conclusiones', level=1)
doc.add_paragraph(
    'El proyecto demuestra de manera práctica cómo las estructuras de datos '
    'nativas de Python se aplican a problemas reales de gestión hospitalaria. '
    'Se logró:'
)
conclusiones = [
    'Modelar un sistema completo usando exclusivamente diccionarios, tuplas, '
    'conjuntos y listas, cada uno justificado por su complejidad algorítmica.',
    'Implementar patrones de flujo de datos (FIFO para colas de espera, '
    'LIFO para control de medicación) usando listas simples.',
    'Simular una lista enlazada con diccionarios como nodos, demostrando '
    'que las estructuras de datos abstractas pueden construirse sobre las '
    'nativas.',
    'Aplicar operaciones de conjuntos a un caso de uso médico real '
    '(cruce de alergias entre pacientes).',
    'Implementar un algoritmo de ordenamiento manual (Burbuja) para '
    'comprender la mecánica interna del ordenamiento.',
    'Diseñar una arquitectura preparada para evolucionar a programación '
    'orientada a objetos en fases posteriores.',
]
for c in conclusiones:
    doc.add_paragraph(c, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph(
    'El sistema queda listo para su refactorización a POO, donde cada '
    'módulo funcional se convertirá en una clase con métodos, manteniendo '
    'la misma lógica de datos pero con encapsulamiento y herencia.'
)

# ── Guardar ───────────────────────────────────────────────────────────────
ruta = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                    'reportes', 'Informe_Proyecto_Sistema_Hospitalario.docx')
os.makedirs(os.path.dirname(ruta), exist_ok=True)
doc.save(ruta)
print(f"Informe generado: {ruta}")
